import logging
import asyncio
import os
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, FSInputFile
import google.generativeai as genai
from docx import Document

# 1. SOZLAMALAR
import os

# Render'dagi o'zgaruvchilardan o'qib olish
BOT_TOKEN = os.getenv("BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
# transport='rest' parametritini majburiy kiritamiz
genai.configure(api_key=GEMINI_API_KEY, transport='rest')

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

class TestStates(StatesGroup):
    waiting_for_content = State()
    waiting_for_count = State()

def main_menu():
    kb = [[KeyboardButton(text="📝 Test yaratish")]]
    return ReplyKeyboardMarkup(keyboard=kb, resize_keyboard=True)

@dp.message(Command("start"))
async def start_cmd(message: types.Message):
    await message.answer("Assalomu alaykum! Test yaratish uchun tugmani bosing.", reply_markup=main_menu())

@dp.message(F.text == "📝 Test yaratish")
async def create_test(message: types.Message, state: FSMContext):
    await message.answer("Mavzuni yozing:")
    await state.set_state(TestStates.waiting_for_content)

@dp.message(TestStates.waiting_for_content)
async def handle_content(message: types.Message, state: FSMContext):
    await state.update_data(topic=message.text)
    await message.answer("Nechta savol kerak?")
    await state.set_state(TestStates.waiting_for_count)

@dp.message(TestStates.waiting_for_count)
async def finalize_test(message: types.Message, state: FSMContext):
    if not message.text.isdigit():
        await message.answer("Son kiriting!")
        return

    data = await state.get_data()
    wait_msg = await message.answer("AI tayyorlamoqda... ⏳")

    try:
        # Modelni chaqirishning eng barqaror usuli
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"O'qituvchi uchun {data['topic']} mavzusida {message.text} ta test va kalitini tuz."
        
        # Javobni olish
        response = model.generate_content(prompt)
        
        # Word fayl yaratish
        doc = Document()
        doc.add_heading('Testlar', 0)
        doc.add_paragraph(response.text)
        
        file_path = f"test_{message.from_user.id}.docx"
        doc.save(file_path)

        await message.answer_document(FSInputFile(file_path), caption="Tayyor! ✅")
        os.remove(file_path)
        
    except Exception as e:
        await message.answer(f"AI bilan bog'lanishda muammo. VPN yoqib ko'ring yoki birozdan so'ng urinib ko'ring.")
        print(f"Xato: {e}")
    
    await wait_msg.delete()
    await state.clear()

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":

    asyncio.run(main())
